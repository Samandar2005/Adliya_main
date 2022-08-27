from docx import Document
from docx.shared import Inches
from django.http import HttpResponse
from django.conf import settings

from home.models import Staff


def form6(request, pk):

    document = Document()

    staff = Staff.objects.get(pk=pk)

    if staff.img:
        document.add_picture(str(settings.BASE_DIR.joinpath(
            *str(staff.img.url).split('/'))), width=Inches(1.25))

    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = staff.last_name
    hdr_cells[1].text = staff.first_name
    hdr_cells[2].text = str(staff.age)
    hdr_cells[3].text = staff.posation.title

    # hdr_cells[4].text = staff.subdepartment.name

    # for qty, id, desc in records:
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = str(qty)
    #     row_cells[1].text = id
    #     row_cells[2].text = desc

    document.add_page_break()

    document.save('demo.docx')

    response = HttpResponse(headers={'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                            'Content-Disposition': 'attachment; filename="download.docx"'})
    document.save(response)

    return response
